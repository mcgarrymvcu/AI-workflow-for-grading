from docx import Document
import tempfile

def generate_feedback_docx(scores, feedback, transcript):
    """
    Generates a .docx feedback file based on grading output.
    :param scores: Dict of rubric item scores
    :param feedback: Summary feedback string
    :param transcript: Transcript text (optional, included for context)
    :return: Path to the .docx file
    """
    doc = Document()

    doc.add_heading("Student Presentation Feedback", level=1)

    doc.add_heading("Score Breakdown", level=2)
    for category, score in scores.items():
        doc.add_paragraph(f"{category}: {score}")

    doc.add_heading("Summary Feedback", level=2)
    doc.add_paragraph(feedback)

    doc.add_heading("Transcript Excerpt", level=2)
    doc.add_paragraph(transcript[:1000] + "...")

    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(output_path)
    return output_path
